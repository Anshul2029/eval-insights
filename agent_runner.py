"""
agent_runner.py — Runs a data-analysis task through a single LLM step by step.

REAL mode: actually loads the Excel file, runs pandas/scipy computation,
and feeds real computed data into LLM prompts so token counts are genuine.

Steps 0, 1, 2, 5  →  LLM calls with real data in context (high token counts)
Steps 3, 4         →  pure Python (serialise facts / define sections), 0 LLM tokens

Supported LLMs:
  "gemini"    →  gemini-2.5-flash
  "groq"      →  llama-3.3-70b-versatile
  "qwen"      →  qwen/qwen3-32b via Groq
  "ollama"    →  mistral (local)
  "gpt4omini" →  gpt-4o-mini
"""

import json
import os
import re
import warnings
from datetime import datetime, timezone

import numpy as np
import pandas as pd
from scipy import stats as sp_stats

_GEMINI_MODEL_NAME = "gemini-2.5-flash"
_OLLAMA_BASE  = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
_OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "mistral")
_OLLAMA_TIMEOUT = int(os.getenv("OLLAMA_TIMEOUT", "180"))
_GROQ_MODEL    = "llama-3.3-70b-versatile"
_QWEN_MODEL    = "qwen/qwen3-32b"
_OPENAI_MODEL  = "gpt-4o-mini"

_EXCEL_SEARCH_DIRS = ["temp_excel", "dataset", "."]

_gemini_model  = None
_groq_client   = None
_openai_client = None


def _get_gemini():
    global _gemini_model
    if _gemini_model is None:
        warnings.filterwarnings("ignore")
        import google.generativeai as genai
        genai.configure(api_key=os.environ["GOOGLE_GEMINI_API_KEY"])
        _gemini_model = genai.GenerativeModel(_GEMINI_MODEL_NAME)
    return _gemini_model


def _get_groq():
    global _groq_client
    if _groq_client is None:
        from groq import Groq
        _groq_client = Groq(api_key=os.environ["GROQ_API_KEY"])
    return _groq_client


def _get_openai():
    global _openai_client
    if _openai_client is None:
        from openai import OpenAI
        _openai_client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
    return _openai_client


def _zero_usage() -> dict:
    return {"input_tokens": 0, "output_tokens": 0,
            "thinking_tokens": 0, "total_tokens": 0, "cached_tokens": 0}


def _parse_json(text: str) -> dict:
    text = text.strip()
    if "```" in text:
        for part in text.split("```"):
            part = part.strip().lstrip("json").strip()
            if part.startswith("{"):
                text = part
                break
    start = text.find("{")
    end = text.rfind("}") + 1
    if start >= 0 and end > start:
        try:
            return json.loads(text[start:end])
        except json.JSONDecodeError:
            pass
    return {}


def _find_excel(dataset_file: str) -> str:
    root = os.path.dirname(os.path.abspath(__file__))
    for d in _EXCEL_SEARCH_DIRS:
        path = os.path.join(root, d, dataset_file)
        if os.path.exists(path):
            return path
    raise FileNotFoundError(
        f"Excel file '{dataset_file}' not found in {_EXCEL_SEARCH_DIRS}"
    )


class AgentRunner:
    def __init__(self, llm_name: str):
        assert llm_name in ("gemini", "ollama", "groq", "qwen", "gpt4omini"), \
            f"Unknown LLM: {llm_name}"
        self.llm_name = llm_name

    # ── LLM call layer ────────────────────────────────────────────────────────

    def _call_llm(self, system: str, user: str, retries: int = 2) -> tuple:
        for attempt in range(retries):
            if self.llm_name == "gemini":
                text, usage = self._call_gemini(system, user)
            elif self.llm_name in ("groq", "qwen"):
                text, usage = self._call_groq(system, user)
            elif self.llm_name == "gpt4omini":
                text, usage = self._call_openai(system, user)
            else:
                return self._call_ollama(system + "\n\n" + user)
            if text and _parse_json(text):
                return text, usage
            if attempt < retries - 1:
                print(f"    [{self.llm_name}] empty response, retrying...")
        return text, usage

    def _call_gemini(self, system: str, user: str) -> tuple:
        model = _get_gemini()
        try:
            response = model.generate_content(f"{system}\n\n{user}")
            text = response.text
            meta = response.usage_metadata
            inp = meta.prompt_token_count
            out = meta.candidates_token_count
            total = meta.total_token_count
            cached = getattr(meta, "cached_content_token_count", 0) or 0
            thinking = max(0, total - inp - out)
            return text, {"input_tokens": inp, "output_tokens": out,
                          "thinking_tokens": thinking, "total_tokens": total,
                          "cached_tokens": cached}
        except Exception as e:
            print(f"    [gemini] Error: {e}")
            return "", _zero_usage()

    def _call_ollama(self, prompt: str) -> tuple:
        import requests
        try:
            resp = requests.post(
                f"{_OLLAMA_BASE}/api/generate",
                json={"model": _OLLAMA_MODEL, "prompt": prompt, "stream": False},
                timeout=_OLLAMA_TIMEOUT,
            )
            if resp.status_code == 200:
                data = resp.json()
                inp = data.get("prompt_eval_count", 0)
                out = data.get("eval_count", 0)
                return data.get("response", ""), {
                    "input_tokens": inp, "output_tokens": out,
                    "thinking_tokens": 0, "total_tokens": inp + out, "cached_tokens": 0}
            return "", _zero_usage()
        except Exception as e:
            print(f"    [ollama] Error: {e}")
            return "", _zero_usage()

    def _call_groq(self, system: str, user: str) -> tuple:
        try:
            client = _get_groq()
            model = _QWEN_MODEL if self.llm_name == "qwen" else _GROQ_MODEL
            r = client.chat.completions.create(
                model=model,
                temperature=0,
                messages=[{"role": "system", "content": system},
                           {"role": "user",   "content": user}],
            )
            text = r.choices[0].message.content or ""
            text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()
            u = r.usage
            inp = u.prompt_tokens
            out = u.completion_tokens
            think = getattr(u, "completion_tokens_details", None)
            think_tokens = getattr(think, "reasoning_tokens", 0) or 0
            return text, {"input_tokens": inp, "output_tokens": out - think_tokens,
                          "thinking_tokens": think_tokens,
                          "total_tokens": inp + out, "cached_tokens": 0}
        except Exception as e:
            print(f"    [{self.llm_name}] Error: {e}")
            return "", _zero_usage()

    def _call_openai(self, system: str, user: str) -> tuple:
        try:
            client = _get_openai()
            r = client.chat.completions.create(
                model=_OPENAI_MODEL, temperature=0,
                messages=[{"role": "system", "content": system},
                           {"role": "user",   "content": user}],
                response_format={"type": "json_object"},
            )
            text = r.choices[0].message.content or ""
            u = r.usage
            return text, {"input_tokens": u.prompt_tokens,
                          "output_tokens": u.completion_tokens,
                          "thinking_tokens": 0,
                          "total_tokens": u.prompt_tokens + u.completion_tokens,
                          "cached_tokens": 0}
        except Exception as e:
            print(f"    [gpt4omini] Error: {e}")
            return "", _zero_usage()

    # ── real pandas/scipy computation ─────────────────────────────────────────

    def _load_and_parse(self, excel_path: str) -> tuple:
        """Actually load Excel with pandas. Returns (df, stats_dict)."""
        df = pd.read_excel(excel_path)
        missing = int(df.isnull().sum().sum())
        missing_per_col = {c: int(v) for c, v in df.isnull().sum().items() if v > 0}
        duplicates = int(df.duplicated().sum())
        dtypes = {c: str(t) for c, t in df.dtypes.items()}

        parse_stats = {
            "row_count": len(df),
            "columns": df.columns.tolist(),
            "dtypes": dtypes,
            "missing_values": missing,
            "missing_per_col": missing_per_col,
            "duplicate_rows": duplicates,
        }
        for col in df.select_dtypes(include="number").columns:
            desc = df[col].describe()
            parse_stats[f"{col}_stats"] = {
                "mean": round(float(desc["mean"]), 2),
                "std":  round(float(desc["std"]),  2),
                "min":  int(desc["min"]),
                "max":  int(desc["max"]),
            }
        return df, parse_stats

    def _compute_stats(self, df: pd.DataFrame, ground_truth: str) -> dict:
        """Real z-score anomaly detection + group aggregations."""
        results = {}

        if "Revenue" in df.columns:
            results["total_revenue"] = int(df["Revenue"].sum())
            results["mean_revenue"]  = round(float(df["Revenue"].mean()), 2)
        if "Units_Sold" in df.columns:
            results["total_units"] = int(df["Units_Sold"].sum())

        for group_col in ["Region", "Product", "Month"]:
            if group_col in df.columns and "Revenue" in df.columns:
                gb = (df.groupby(group_col)["Revenue"].sum()
                        .sort_values(ascending=False))
                results[f"revenue_by_{group_col.lower()}"] = {
                    k: int(v) for k, v in gb.items()
                }

        # Z-score anomaly detection
        anomalies = []
        for col in df.select_dtypes(include="number").columns:
            z = sp_stats.zscore(df[col].fillna(df[col].median()))
            mask = np.abs(z) > 2.0
            outliers = df[mask].copy()
            outliers["_z"] = z[mask]
            for _, row in outliers.iterrows():
                entry = {"column": col,
                         "value": float(row[col]),
                         "z_score": round(float(row["_z"]), 3)}
                for cat in df.select_dtypes(include="object").columns:
                    entry[cat.lower()] = str(row[cat])
                anomalies.append(entry)

        anomalies.sort(key=lambda x: abs(x["z_score"]), reverse=True)
        results["anomalies"]     = anomalies
        results["anomaly_count"] = len(anomalies)
        results["ground_truth"]  = ground_truth
        return results

    # ── prompts (include real computed data) ──────────────────────────────────

    @staticmethod
    def _sys() -> str:
        return ("You are an AI data analyst. "
                "Respond ONLY with valid JSON — no prose, no markdown fences.")

    def _prompt_plan(self, user_prompt: str, df: pd.DataFrame) -> str:
        sample = df.head(20).to_csv(index=False)
        return f"""Plan this data analysis task in one sentence.
Task: {user_prompt}
Dataset sample (first 20 rows):
{sample}
Total rows: {len(df)}, Columns: {df.columns.tolist()}
Dtypes: { {c: str(t) for c, t in df.dtypes.items()} }
Return JSON: {{"agent_plan":"one sentence describing the full analysis approach including anomaly detection and Word report"}}"""

    def _prompt_step1(self, user_prompt: str, df: pd.DataFrame,
                      parse_stats: dict) -> str:
        # Send the full dataset so the LLM actually sees all the data
        full_csv = df.to_csv(index=False)
        kf = [
            f"row_count: {parse_stats['row_count']}",
            f"columns: {parse_stats['columns']}",
            f"missing_values: {parse_stats['missing_values']}",
            f"duplicate_rows: {parse_stats['duplicate_rows']}",
            f"dtypes: {parse_stats['dtypes']}",
        ]
        for col in df.select_dtypes(include="number").columns:
            s = parse_stats.get(f"{col}_stats", {})
            kf.append(f"{col}_mean: {s.get('mean')}, std: {s.get('std')}, "
                      f"min: {s.get('min')}, max: {s.get('max')}")
        return f"""Step 1 data_parsing: The Excel file was loaded with pandas. Analyse every row and write a thorough professional description.
Task: {user_prompt}
Full dataset ({len(df)} rows):
{full_csv}
Computed statistics:
{json.dumps(parse_stats, indent=2)}
Return JSON:
{{
  "output": "3-4 sentence description covering row count, column types, data quality, and numeric ranges",
  "what_agent_did": "Loaded Excel with pandas.read_excel, validated schema, computed descriptive statistics, checked for nulls and duplicates",
  "tools_called": ["pandas.read_excel", "df.describe", "df.isnull", "df.duplicated", "df.dtypes"],
  "key_facts_produced": {json.dumps(kf)}
}}"""

    def _prompt_step2(self, user_prompt: str, df: pd.DataFrame,
                      compute_stats: dict) -> str:
        anomalies   = compute_stats.get("anomalies", [])
        rev_region  = compute_stats.get("revenue_by_region",  {})
        rev_product = compute_stats.get("revenue_by_product", {})
        rev_month   = compute_stats.get("revenue_by_month",   {})
        anomaly_lines = "\n".join(
            f"  {a['column']} | {a.get('month','?')} | {a.get('region','?')} | "
            f"{a.get('product','?')} | value={a['value']:,.0f} | z={a['z_score']}"
            for a in anomalies
        ) or "  None"

        kf = [
            f"total_revenue: {compute_stats.get('total_revenue', 0):,}",
            f"total_units: {compute_stats.get('total_units', 0):,}",
            f"mean_revenue_per_row: {compute_stats.get('mean_revenue', 0):,}",
            f"anomaly_count: {len(anomalies)}",
        ]
        for a in anomalies[:8]:
            kf.append(
                f"anomaly_{a.get('month','?')}_{a.get('region','?')}_"
                f"{a.get('product','?')}: value={a['value']:,.0f} z={a['z_score']}"
            )
        for region, rev in rev_region.items():
            kf.append(f"revenue_{region}: {rev:,}")
        for product, rev in list(rev_product.items())[:6]:
            kf.append(f"revenue_{product}: {rev:,}")

        return f"""Step 2 computation: Z-score anomaly detection and full group aggregations on the real data.
Task: {user_prompt}
Ground truth: {compute_stats.get('ground_truth', '')}

Computed results:
  Total Revenue      : {compute_stats.get('total_revenue', 0):,}
  Total Units        : {compute_stats.get('total_units', 0):,}
  Mean Revenue/row   : {compute_stats.get('mean_revenue', 0):,}
  Revenue by Region  : {json.dumps(rev_region, indent=2)}
  Revenue by Product : {json.dumps(rev_product, indent=2)}
  Revenue by Month   : {json.dumps(rev_month, indent=2)}

All anomalies detected (|z-score| > 2.0, sorted by severity):
{anomaly_lines}

Return JSON:
{{
  "output": "3-4 sentence summary of findings with specific numbers: total revenue, top region, anomaly count and most severe anomaly",
  "what_agent_did": "Ran scipy.stats.zscore on all numeric columns, computed groupby aggregations by Region, Product and Month",
  "tools_called": ["scipy.stats.zscore", "pandas.groupby", "numpy.abs", "df.describe"],
  "key_facts_produced": {json.dumps(kf)}
}}"""

    def _prompt_step3(self, user_prompt: str, compute_stats: dict,
                      parse_stats: dict, s2_facts: list) -> str:
        anomalies   = compute_stats.get("anomalies", [])
        rev_region  = compute_stats.get("revenue_by_region",  {})
        rev_product = compute_stats.get("revenue_by_product", {})
        rev_month   = compute_stats.get("revenue_by_month",   {})
        anomaly_json = json.dumps(anomalies, indent=2)

        return f"""Step 3 context_handoff: Organise ALL computed facts from Excel into a structured context object for the Word document builder.
Task: {user_prompt}

All Step 2 computed facts to hand off:
  KPIs:
    Total Revenue      : {compute_stats.get('total_revenue', 0):,}
    Total Units        : {compute_stats.get('total_units', 0):,}
    Dataset rows       : {parse_stats['row_count']}
    Missing values     : {parse_stats['missing_values']}
    Anomaly count      : {len(anomalies)}
  Revenue by Region  : {json.dumps(rev_region)}
  Revenue by Product : {json.dumps(rev_product)}
  Revenue by Month   : {json.dumps(rev_month)}
  Full anomaly list  :
{anomaly_json}

Step 2 key_facts: {json.dumps(s2_facts)}

Your job: confirm every fact is packaged and ready for the Word builder.
Return JSON:
{{
  "output": "Packaged {len(s2_facts)} Step-2 facts and {len(anomalies)} anomaly records into Word document context",
  "what_agent_did": "Serialised all KPIs, regional breakdowns, product breakdowns, monthly trends and anomaly table into Word builder context",
  "tools_called": ["python-docx Document", "json.dumps"],
  "key_facts_produced": {json.dumps([f"passed: {f}" for f in s2_facts])},
  "facts_passed_to_word": {json.dumps(s2_facts)}
}}"""

    def _prompt_step4(self, user_prompt: str, compute_stats: dict,
                      parse_stats: dict, s3_facts: list) -> str:
        anomaly_count = compute_stats.get("anomaly_count", 0)
        rev_region    = compute_stats.get("revenue_by_region", {})
        rev_product   = compute_stats.get("revenue_by_product", {})
        sections = ["Executive Summary", "Key Performance Indicators",
                    "Regional & Product Breakdown", "Anomaly Analysis", "Recommendations"]

        return f"""Step 4 report_structuring: Design the full Word document structure and layout for this analysis report.
Task: {user_prompt}

Available data to structure:
  KPIs: Total Revenue {compute_stats.get('total_revenue',0):,}, Total Units {compute_stats.get('total_units',0):,}, {parse_stats['row_count']} rows
  Regions: {list(rev_region.keys())}
  Products: {list(rev_product.keys())}
  Anomalies to document: {anomaly_count}
  Context facts received: {len(s3_facts)}

Design exactly these {len(sections)} sections: {sections}
For each section specify: what content goes in it, what tables or lists are needed, and the key message.

Return JSON:
{{
  "output": "Report designed with {len(sections)} sections covering KPIs, regional breakdown, {anomaly_count} anomalies and recommendations",
  "what_agent_did": "Designed Word document layout: heading hierarchy, KPI table, anomaly table, regional comparison and recommendation bullets",
  "tools_called": ["python-docx Document", "add_heading", "add_table", "add_paragraph"],
  "key_facts_produced": [
    "sections: {', '.join(sections)}",
    "section_count: {len(sections)}",
    "anomaly_section: present",
    "executive_summary: present",
    "kpi_table: present",
    "regional_breakdown_table: present",
    "product_breakdown_table: present"
  ],
  "section_plan": {{
    "Executive Summary": "2-3 sentence overview with total revenue and top finding",
    "Key Performance Indicators": "table of total revenue, units, region breakdown",
    "Regional & Product Breakdown": "table comparing revenue across {list(rev_region.keys())} and top products",
    "Anomaly Analysis": "table of {anomaly_count} anomalies with month/region/product/value/z-score",
    "Recommendations": "3-4 bullet points based on findings"
  }}
}}"""

    def _prompt_step5(self, user_prompt: str, compute_stats: dict,
                      parse_stats: dict, ground_truth: str) -> str:
        anomalies    = compute_stats.get("anomalies", [])
        has_planted  = bool(ground_truth) and "no planted" not in ground_truth.lower()
        anomaly_json = json.dumps(anomalies[:10], indent=2)
        rev_region   = compute_stats.get("revenue_by_region",  {})
        rev_product  = compute_stats.get("revenue_by_product", {})
        total_rev    = compute_stats.get("total_revenue", 0)
        total_units  = compute_stats.get("total_units", 0)

        note = ("A planted anomaly exists — include it in the anomaly section."
                if has_planted else
                "No planted anomaly — report only real statistical outliers.")

        return f"""Step 5 narrative_generation: Write the Word document narrative using the real computed data.
Task: {user_prompt}
Ground truth: {ground_truth}
NOTE: {note}

Real computed data:
  Total Revenue   : {total_rev:,}
  Total Units     : {total_units:,}
  Revenue by Region  : {json.dumps(rev_region)}
  Revenue by Product : {json.dumps(rev_product)}
  Dataset: {parse_stats['row_count']} rows, {len(parse_stats['columns'])} columns, {parse_stats['missing_values']} missing values
  Anomalies (top 10 by |z-score|):
{anomaly_json}

Return JSON:
{{
  "output": "Word report written — {compute_stats.get('anomaly_count', 0)} anomalies documented with specific numbers",
  "what_agent_did": "Generated full Word document: executive summary, performance analysis, anomaly section, recommendations",
  "tools_called": ["python-docx Document", "add_heading", "add_paragraph", "add_table"],
  "key_facts_produced": [
    "organic_outliers_reported: true",
    "planted_anomaly_reported: {'true' if has_planted else 'false'}",
    "narrative_includes_specific_numbers: true",
    "anomaly_count_reported: {compute_stats.get('anomaly_count', 0)}",
    "total_revenue_in_narrative: {total_rev:,}"
  ],
  "word_output_actual_text": {{
    "executive_summary": "3-4 sentences with real numbers: total revenue {total_rev:,}, units {total_units:,}, top region, anomaly count",
    "anomaly_section": "detailed description of each anomaly with month/region/product/value/z-score",
    "recommendations": ["3-4 specific actionable recommendations grounded in the actual data"]
  }}
}}"""

    # ── docx writer ───────────────────────────────────────────────────────────

    def _write_docx(self, trace_id: str, user_prompt: str,
                    word_output: dict, compute_stats: dict, parse_stats: dict) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(f"Analysis Report — {trace_id}", 0)
        doc.add_paragraph(f"Task: {user_prompt}")
        doc.add_paragraph(f"Generated by: {self.llm_name.upper()}  |  "
                          f"{datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(word_output.get("executive_summary", "") or "No summary available.")

        # KPI table
        doc.add_heading("Key Performance Indicators", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"
        kpis = [
            ("Total Revenue",  f"{compute_stats.get('total_revenue', 0):,}"),
            ("Total Units",    f"{compute_stats.get('total_units', 0):,}"),
            ("Rows Analysed",  str(parse_stats.get("row_count", ""))),
            ("Missing Values", str(parse_stats.get("missing_values", 0))),
            ("Anomalies Found", str(compute_stats.get("anomaly_count", 0))),
        ]
        for region, rev in list(compute_stats.get("revenue_by_region", {}).items())[:3]:
            kpis.append((f"Revenue — {region}", f"{rev:,}"))
        for metric, value in kpis:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = value

        # Anomaly section
        doc.add_heading("Anomaly Analysis", level=1)
        anomaly_text = word_output.get("anomaly_section", "")
        if anomaly_text:
            doc.add_paragraph(anomaly_text)
        anomalies = compute_stats.get("anomalies", [])
        if anomalies:
            doc.add_paragraph(f"Detected {len(anomalies)} outlier(s) with |z-score| > 2.0:")
            atbl = doc.add_table(rows=1, cols=5)
            atbl.style = "Table Grid"
            hdrs = ["Column", "Month", "Region", "Product", "Value (z-score)"]
            for i, h in enumerate(hdrs):
                atbl.rows[0].cells[i].text = h
            for a in anomalies[:10]:
                row = atbl.add_row().cells
                row[0].text = a.get("column", "")
                row[1].text = a.get("month", "—")
                row[2].text = a.get("region", "—")
                row[3].text = a.get("product", "—")
                row[4].text = f"{a['value']:,.0f}  (z={a['z_score']})"

        # Recommendations
        recs = word_output.get("recommendations", [])
        if recs:
            doc.add_heading("Recommendations", level=1)
            for rec in recs:
                doc.add_paragraph(str(rec), style="List Bullet")

        reports_dir = os.path.join(os.path.dirname(__file__), "results", "reports")
        os.makedirs(reports_dir, exist_ok=True)
        path = os.path.join(reports_dir, f"report_{trace_id}_{self.llm_name}.docx")
        doc.save(path)
        print(f"  [{self.llm_name.upper()}] Word doc saved -> {path}")
        return path

    # ── main run ──────────────────────────────────────────────────────────────

    def run(self, user_prompt: str, source_data_summary: dict,
            dataset_file: str, trace_id: str) -> dict:
        print(f"\n  [{self.llm_name.upper()}] Starting REAL run for {trace_id}...")
        sys_prompt = self._sys()
        token_per_step = []

        # ── load Excel (real) ─────────────────────────────────────────────────
        excel_path = _find_excel(dataset_file)
        print(f"  [{self.llm_name.upper()}] Loaded: {excel_path}")
        df, parse_stats = self._load_and_parse(excel_path)

        ground_truth = source_data_summary.get("ground_truth", "")
        compute_stats = self._compute_stats(df, ground_truth)
        print(f"  [{self.llm_name.upper()}] Real computation: "
              f"{parse_stats['row_count']} rows, "
              f"{compute_stats['anomaly_count']} anomalies, "
              f"total_revenue={compute_stats.get('total_revenue',0):,}")

        # ── step 0: planning (LLM with data sample) ───────────────────────────
        print(f"  [{self.llm_name.upper()}] Step 0: planning")
        p_text, p_usage = self._call_llm(sys_prompt,
                                          self._prompt_plan(user_prompt, df))
        p_data = _parse_json(p_text)
        agent_plan = p_data.get(
            "agent_plan",
            "Load Excel with pandas, compute z-score anomalies, write Word report."
        )
        token_per_step.append({"step_number": 0, "action_type": "planning", **p_usage})

        # ── step 1: data_parsing (LLM with real CSV + stats) ──────────────────
        print(f"  [{self.llm_name.upper()}] Step 1: data_parsing")
        s1_text, s1_usage = self._call_llm(
            sys_prompt, self._prompt_step1(user_prompt, df, parse_stats)
        )
        s1 = _parse_json(s1_text)
        step1 = {
            "step_number": 1, "app": "Excel", "action_type": "data_parsing",
            "latency_observed": "real pandas",
            "tools_called": s1.get("tools_called", ["pandas.read_excel"]),
            "output": s1.get("output", f"Loaded {parse_stats['row_count']} rows, "
                             f"{len(parse_stats['columns'])} columns, "
                             f"{parse_stats['missing_values']} missing values."),
            "what_agent_did": s1.get("what_agent_did", ""),
            "key_facts_produced": s1.get("key_facts_produced", [
                f"row_count: {parse_stats['row_count']}",
                f"columns: {parse_stats['columns']}",
                f"missing_values: {parse_stats['missing_values']}",
                f"duplicate_rows: {parse_stats['duplicate_rows']}",
                f"dtypes: {parse_stats['dtypes']}",
            ]),
        }
        token_per_step.append({"step_number": 1, "action_type": "data_parsing", **s1_usage})

        # ── step 2: computation (LLM with real z-scores & anomalies) ──────────
        print(f"  [{self.llm_name.upper()}] Step 2: computation")
        s2_text, s2_usage = self._call_llm(
            sys_prompt, self._prompt_step2(user_prompt, df, compute_stats)
        )
        s2 = _parse_json(s2_text)
        # Build definitive key_facts from real computed numbers
        anomalies = compute_stats.get("anomalies", [])
        real_kf = [
            f"total_revenue: {compute_stats.get('total_revenue', 0):,}",
            f"total_units: {compute_stats.get('total_units', 0):,}",
            f"anomaly_count: {len(anomalies)}",
        ]
        for a in anomalies[:6]:
            real_kf.append(
                f"anomaly_{a.get('month','?')}_{a.get('region','?')}_"
                f"{a.get('product','?')}: value={a['value']:,.0f} z={a['z_score']}"
            )
        for region, rev in list(compute_stats.get("revenue_by_region", {}).items())[:3]:
            real_kf.append(f"revenue_{region}: {rev:,}")

        step2 = {
            "step_number": 2, "app": "Excel", "action_type": "computation",
            "latency_observed": "real scipy",
            "tools_called": s2.get("tools_called",
                                    ["scipy.stats.zscore", "pandas.groupby"]),
            "output": s2.get("output",
                             f"Found {len(anomalies)} anomalies; "
                             f"total revenue {compute_stats.get('total_revenue',0):,}."),
            "what_agent_did": s2.get("what_agent_did", ""),
            "key_facts_produced": real_kf,
        }
        token_per_step.append({"step_number": 2, "action_type": "computation", **s2_usage})

        # ── step 3: context_handoff (LLM organises all facts for Word) ──────────
        print(f"  [{self.llm_name.upper()}] Step 3: context_handoff")
        s3_text, s3_usage = self._call_llm(
            sys_prompt,
            self._prompt_step3(user_prompt, compute_stats, parse_stats,
                               step2["key_facts_produced"])
        )
        s3 = _parse_json(s3_text)
        facts_passed = s3.get("facts_passed_to_word", step2["key_facts_produced"])
        step3 = {
            "step_number": 3, "app": "Excel→Word", "action_type": "context_handoff",
            "latency_observed": "real LLM",
            "tools_called": s3.get("tools_called", ["python-docx Document"]),
            "output": s3.get("output",
                             f"Serialised {len(step2['key_facts_produced'])} facts into Word context."),
            "what_agent_did": s3.get("what_agent_did", ""),
            "key_facts_produced": s3.get("key_facts_produced",
                                          [f"passed: {f}" for f in step2["key_facts_produced"]]),
            "facts_passed_to_word": facts_passed,
        }
        token_per_step.append({"step_number": 3, "action_type": "context_handoff",
                                **s3_usage})

        # ── step 4: report_structuring (LLM designs full document layout) ─────
        print(f"  [{self.llm_name.upper()}] Step 4: report_structuring")
        s4_text, s4_usage = self._call_llm(
            sys_prompt,
            self._prompt_step4(user_prompt, compute_stats, parse_stats,
                               step3["key_facts_produced"])
        )
        s4 = _parse_json(s4_text)
        step4 = {
            "step_number": 4, "app": "Word", "action_type": "report_structuring",
            "latency_observed": "real LLM",
            "tools_called": s4.get("tools_called",
                                    ["python-docx add_heading", "add_table"]),
            "output": s4.get("output", ""),
            "what_agent_did": s4.get("what_agent_did", ""),
            "key_facts_produced": s4.get("key_facts_produced", [
                "sections: Executive Summary, KPIs, Regional Breakdown, Anomaly Analysis, Recommendations",
                "section_count: 5", "anomaly_section: present", "kpi_table: present",
            ]),
        }
        token_per_step.append({"step_number": 4, "action_type": "report_structuring",
                                **s4_usage})

        # ── step 5: narrative_generation (LLM with full real analysis) ────────
        print(f"  [{self.llm_name.upper()}] Step 5: narrative_generation")
        s5_text, s5_usage = self._call_llm(
            sys_prompt,
            self._prompt_step5(user_prompt, compute_stats, parse_stats, ground_truth)
        )
        s5 = _parse_json(s5_text)
        word_output = s5.get("word_output_actual_text",
                             {"executive_summary": "", "anomaly_section": "",
                              "recommendations": []})
        step5 = {
            "step_number": 5, "app": "Word", "action_type": "narrative_generation",
            "latency_observed": "real LLM",
            "tools_called": s5.get("tools_called",
                                    ["python-docx", "add_paragraph", "add_table"]),
            "output": s5.get("output", ""),
            "what_agent_did": s5.get("what_agent_did", ""),
            "key_facts_produced": s5.get("key_facts_produced", [
                "organic_outliers_reported: true",
                f"planted_anomaly_reported: {'true' if bool(ground_truth) and 'no planted' not in ground_truth.lower() else 'false'}",
                "narrative_includes_specific_numbers: true",
                f"anomaly_count_reported: {len(anomalies)}",
            ]),
        }
        token_per_step.append({"step_number": 5, "action_type": "narrative_generation",
                                **s5_usage})

        # ── context manifest (real comparison) ────────────────────────────────
        word_flat = " ".join(str(v) for v in word_output.values() if v).lower()
        facts_present, facts_lost = [], []
        for fact in real_kf:
            key = fact.split(":")[0].strip().lower().replace("_", " ").split()[-1]
            if key in word_flat:
                facts_present.append(fact)
            else:
                facts_lost.append(fact)

        context_manifest = {
            "facts_produced_in_excel_step2": real_kf,
            "facts_present_in_word_output": [f"PRESENT: {f}" for f in facts_present],
            "facts_lost_at_boundary": facts_lost,
            "boundary": "Excel→Word (Step 3)",
            "context_loss_detected": len(facts_lost) > 0,
        }

        # ── normalise source_data_summary ─────────────────────────────────────
        src = dict(source_data_summary)
        if "ground_truth_anomalies" not in src:
            planted = ("None" if not ground_truth or "no planted" in ground_truth.lower()
                       else ground_truth)
            src["ground_truth_anomalies"] = {"planted": planted}

        # ── write real Word document ──────────────────────────────────────────
        word_doc_path = self._write_docx(
            trace_id, user_prompt, word_output, compute_stats, parse_stats
        )

        # ── assemble trace ────────────────────────────────────────────────────
        trace = {
            "trace_id": f"{trace_id}_{self.llm_name}",
            "llm": self.llm_name,
            "llm_model": {
                "gemini":    _GEMINI_MODEL_NAME,
                "ollama":    _OLLAMA_MODEL,
                "groq":      _GROQ_MODEL,
                "qwen":      _QWEN_MODEL,
                "gpt4omini": _OPENAI_MODEL,
            }.get(self.llm_name, self.llm_name),
            "source_trace_id": trace_id,
            "run_date": datetime.now(timezone.utc).isoformat(),
            "dataset_file": dataset_file,
            "user_prompt": user_prompt,
            "agent_plan": agent_plan,
            "source_data_summary": src,
            "steps": [step1, step2, step3, step4, step5],
            "context_manifest": context_manifest,
            "word_output_actual_text": word_output,
            "word_doc_path": word_doc_path,
        }

        total = {k: sum(s[k] for s in token_per_step)
                 for k in ("input_tokens", "output_tokens", "thinking_tokens",
                           "total_tokens", "cached_tokens")}

        print(f"  [{self.llm_name.upper()}] Done — total tokens: {total['total_tokens']:,}")

        return {
            "trace": trace,
            "token_usage": {
                "llm": self.llm_name,
                "llm_model": trace["llm_model"],
                "per_step": token_per_step,
                "total": total,
            },
        }

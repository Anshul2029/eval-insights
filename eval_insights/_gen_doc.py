"""Generate PLATFORM_OVERVIEW.docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

NAVY = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xDC, 0x26, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Title ──
title = doc.add_heading("Eval Insights Platform", level=0)
title.runs[0].font.color.rgb = NAVY
title.runs[0].font.size = Pt(22)
sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(12)
r = sub.add_run("Trajectory-Level Failure Analysis for Excel Copilot Eval")
r.font.size = Pt(12)
r.font.color.rgb = MUTED
r.font.italic = True

# ── GOAL ──
h = doc.add_heading("Goal (WHY)", level=1)
h.runs[0].font.color.rgb = NAVY

p = doc.add_paragraph()
p.add_run("Excel Copilot eval runs produce pass/fail verdicts — but ").font.color.rgb = DARK
p.add_run("not why").font.bold = True
p.add_run(" something failed, ").font.color.rgb = DARK
p.add_run("how").font.bold = True
p.add_run(" the agent behaved during execution, or ").font.color.rgb = DARK
p.add_run("what patterns").font.bold = True
p.add_run(" repeat across hundreds of queries. Teams reviewing eval results were manually reading JSON trajectories or relying on aggregate pass rates that hide critical failure modes.").font.color.rgb = DARK

p2 = doc.add_paragraph()
r2 = p2.add_run("The goal: ")
r2.bold = True
r2.font.color.rgb = NAVY
p2.add_run("Turn raw eval output into actionable intelligence — so the team knows exactly what to fix, what's working, and where effort will have the highest ROI — without reading a single trajectory manually.").font.color.rgb = DARK

# ── APPROACH ──
h = doc.add_heading("Approach (HOW)", level=1)
h.runs[0].font.color.rgb = NAVY

steps = [
    ("1. Trajectory Parsing & Signal Extraction",
     "Parse raw playOutput JSON files — extract every agent step (user query, script execution, script response, agent message) into a structured timeline. Compute trajectory-level signals: step count, error positions, error types, execution time, script similarity groups (retry loops), success claims vs actual results."),
    ("2. LLM-Powered Failure Classification (Two-Axis)",
     "Trajectory Axis — where in the agent's process it went wrong (e.g., Misguided Strategy, Failed Error Recovery, Premature Completion). Outcome Axis — what's wrong with the final output (e.g., Wrong Values, Incomplete Output). Each failure gets a root cause explanation and suggested fix, not just a category label."),
    ("3. Multi-Dimensional Analysis",
     "Success Analysis: Golden trajectories, recovery patterns, passing operations. Failure Analysis: Error-triggering operations, actual error messages, step efficiency gap, wasted steps, retry loop detection. Root Cause Analysis: LLM synthesizes failures into grouped root causes. False Claim Detection: Queries where the agent claims success but the grader disagrees. Recovery Analysis: Cross-references recoverable vs persistent errors."),
    ("4. Cross-Run Evolution Tracking",
     "Upload multiple runs → track per-query status changes (Pass→Fail, Fail→Pass, Consistent). Categorize query evolution patterns with LLM analysis — why regressions happened, what improved. Trajectory diff: how the agent's step-by-step behavior changes across runs for the same query."),
    ("5. Exportable Reports",
     "Self-contained HTML report covering all analysis tabs — shareable offline with stakeholders. Includes: Executive Summary, Dashboard, Deep Analysis, and per-query detail with categories and root causes."),
]

for title_text, body_text in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    pb = doc.add_paragraph(body_text)
    pb.paragraph_format.left_indent = Inches(0.25)
    for run in pb.runs:
        run.font.color.rgb = DARK
        run.font.size = Pt(10)

# ── IMPACT ──
h = doc.add_heading("Impact (WHAT Gains)", level=1)
h.runs[0].font.color.rgb = NAVY

sh = doc.add_heading("Previously Not Possible → Now Unlocked", level=2)
sh.runs[0].font.color.rgb = NAVY
sh.runs[0].font.size = Pt(12)

rows_data = [
    ("Why a query failed", "Manual trajectory reading", "Automated two-axis classification with root cause"),
    ("Failure pattern grouping", "Ad-hoc, per-person", "Systematic: trajectory + outcome categories across all failures"),
    ("Recovery intelligence", "Not tracked", "Identifies which errors are recoverable vs persistent"),
    ("False confidence detection", "Hidden in pass/fail", "Flags queries where agent claims success incorrectly"),
    ("Wasted effort quantification", "Unknown", "Counts exact steps wasted after first error — maps to latency/cost savings"),
    ("Retry loop detection", "Not visible", "Identifies queries where agent repeats near-identical failing scripts"),
    ("Cross-run regression tracking", "Manual diff of reports", "Automated per-query status change tracking with trajectory diff"),
    ("Intent-level performance", "Aggregate pass rate only", "Pass rate by query intent, specificity, complexity"),
    ("Actionable recommendations", 'Generic "improve error handling"', "LLM-generated, specific to the data: names error types, operations, patterns"),
    ("Shareable analysis", "Screenshots or raw JSON", "One-click HTML report with full analysis"),
]

table = doc.add_table(rows=len(rows_data) + 1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"

headers = ["Capability", "Before", "Now"]
for i, hdr in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    r = cell.paragraphs[0].add_run(hdr)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9.5)
    shading = cell._element.get_or_add_tcPr()
    bg = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "1F4E79",
    })
    shading.append(bg)

for row_idx, (cap, before, now) in enumerate(rows_data):
    row = table.rows[row_idx + 1]
    row.cells[0].text = ""
    r = row.cells[0].paragraphs[0].add_run(cap)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK

    row.cells[1].text = ""
    r = row.cells[1].paragraphs[0].add_run(before)
    r.font.size = Pt(9)
    r.font.color.rgb = RED

    row.cells[2].text = ""
    r = row.cells[2].paragraphs[0].add_run(now)
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(1.8)
    row.cells[2].width = Inches(2.8)

# ── Key Insights ──
doc.add_paragraph()
sh2 = doc.add_heading("Key Insights Surfaced", level=2)
sh2.runs[0].font.color.rgb = NAVY
sh2.runs[0].font.size = Pt(12)

insights = [
    ("Success Analysis", "reveals golden trajectories cluster in specific batches and intent types — these are the agent's strength zones and optimization benchmarks"),
    ("Failure Analysis", "shows certain operations (e.g., conditional formatting, multi-sheet references) trigger disproportionate errors — prioritizable fix targets"),
    ("Root Cause Analysis", 'groups failures by underlying cause — e.g., "agent miscalculates because it doesn\'t handle null cells" vs "agent targets wrong sheet"'),
    ("Recovery vs Failure cross-reference", "identifies errors the agent can recover from — the recovery strategy can be extracted and applied to currently-failing queries"),
    ("False claim rate", "quantifies the agent's self-evaluation blind spot — a direct input for adding verification steps"),
    ("Wasted steps metric", "translates directly to potential latency/cost reduction — X steps could be saved if the agent stopped after first error"),
]

for label, desc in insights:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.color.rgb = DARK
    r2.font.size = Pt(10)

# ── Footer ──
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Platform: Streamlit + LLM (Groq/Gemini)  |  Input: evalVNext playOutput JSON  |  Output: Interactive dashboard + downloadable HTML report")
r.font.size = Pt(8.5)
r.font.color.rgb = MUTED
r.italic = True

out = r"c:\Users\t-ashende\Documents\evaluator\eval_insights\Eval_Insights_Platform_Overview.docx"
doc.save(out)
print(f"Saved: {out}")
